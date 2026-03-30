"""
岩土工程写作系统 - Word 导出器
Geo Engineering Writer - Word Exporter

功能：
- 将 Markdown 报告转换为 Word 格式
- 支持加载公司标准模板
- 自动设置岩土报告标准样式
- 表格、标题、段落格式化
"""

from pathlib import Path
from typing import Optional, List, Dict, Any
import re


class GeoWordExporter:
    """岩土工程报告 Word 导出器"""
    
    def __init__(self, template_path: Optional[str] = None):
        """
        初始化 Word 导出器
        
        参数：
            template_path: Word 模板文件路径（可选）
                           如果提供，将基于模板创建文档
        """
        try:
            from docx import Document
            from docx.shared import Pt, Cm, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            
            self.Document = Document
            self.Pt = Pt
            self.Cm = Cm
            self.Inches = Inches
            self.WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH
            self.qn = qn
            
            # 加载模板或创建新文档
            if template_path and Path(template_path).exists():
                self.doc = Document(template_path)
                self.template_path = template_path
            else:
                self.doc = Document()
                self.template_path = None
            
            # 设置标准样式
            self._setup_styles()
            
        except ImportError:
            raise ImportError(
                "请先安装 python-docx: py -m pip install python-docx\n"
                "或者: pip install python-docx"
            )
    
    def _setup_styles(self):
        """设置岩土报告标准样式"""
        try:
            # 设置默认字体为宋体
            style = self.doc.styles["Normal"]
            font = style.font
            font.name = "宋体"
            font.size = self.Pt(12)
            # 中文字体设置
            font._element.rPr.rFonts.set(self.qn('w:eastAsia'), '宋体')
            
            # 设置段落格式
            paragraph_format = style.paragraph_format
            paragraph_format.line_spacing = 1.5  # 1.5倍行距
            paragraph_format.space_before = self.Pt(6)
            paragraph_format.space_after = self.Pt(6)
            
            # 设置标题样式
            for level in range(1, 4):
                heading_style = self.doc.styles[f"Heading {level}"]
                heading_font = heading_style.font
                heading_font.name = "黑体"
                heading_font.size = self.Pt(16 - level * 2)
                heading_font.bold = True
                heading_font._element.rPr.rFonts.set(self.qn('w:eastAsia'), '黑体')
                
                heading_para = heading_style.paragraph_format
                heading_para.space_before = self.Pt(12)
                heading_para.space_after = self.Pt(6)
            
        except Exception as e:
            print(f"样式设置警告: {e}")
    
    def add_section(self, title: str, content: str = "", level: int = 1):
        """
        添加章节
        
        参数：
            title: 章节标题
            content: 章节内容（可选）
            level: 标题级别（1-3）
        """
        # 添加标题
        heading = self.doc.add_heading(title, level=level)
        heading.alignment = self.WD_ALIGN_PARAGRAPH.LEFT
        
        # 添加内容
        if content and content.strip():
            para = self.doc.add_paragraph(content)
            para.alignment = self.WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def add_paragraph(self, text: str, style: str = "Normal"):
        """
        添加段落
        
        参数：
            text: 段落文本
            style: 段落样式
        """
        para = self.doc.add_paragraph(text, style=style)
        return para
    
    def add_bullet_list(self, items: List[str]):
        """
        添加无序列表
        
        参数：
            items: 列表项列表
        """
        for item in items:
            self.doc.add_paragraph(item, style="List Bullet")
    
    def add_numbered_list(self, items: List[str]):
        """
        添加有序列表
        
        参数：
            items: 列表项列表
        """
        for item in items:
            self.doc.add_paragraph(item, style="List Number")
    
    def add_table_from_dict(self, data: List[Dict[str, Any]], headers: List[str]):
        """
        从字典列表添加表格
        
        参数：
            data: 数据列表（每个元素是一个字典）
            headers: 表头列表
        """
        if not data or not headers:
            return
        
        # 创建表格
        table = self.doc.add_table(rows=len(data) + 1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        
        # 写入表头
        for j, header in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = str(header)
            # 设置表头加粗
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
        
        # 写入数据
        for i, row_data in enumerate(data):
            for j, header in enumerate(headers):
                cell = table.rows[i + 1].cells[j]
                value = row_data.get(header, "")
                cell.text = str(value) if value else ""
    
    def add_table_from_markdown(self, markdown_table: str):
        """
        从 Markdown 表格添加表格
        
        参数：
            markdown_table: Markdown 格式的表格字符串
        """
        try:
            # 解析 Markdown 表格
            lines = [line.strip() for line in markdown_table.strip().split('\n') if line.strip()]
            
            if len(lines) < 2:
                return
            
            # 提取表头
            headers = [cell.strip() for cell in lines[0].split('|')[1:-1]]
            
            # 提取数据行（跳过分隔行）
            data_rows = []
            for line in lines[2:]:
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if cells:
                    data_rows.append({headers[i]: cells[i] for i in range(min(len(headers), len(cells)))})
            
            # 添加表格
            self.add_table_from_dict(data_rows, headers)
            
        except Exception as e:
            print(f"Markdown 表格解析失败: {e}")
    
    def add_page_break(self):
        """添加分页符"""
        self.doc.add_page_break()
    
    def add_horizontal_rule(self):
        """添加水平线"""
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = self.WD_ALIGN_PARAGRAPH.CENTER
    
    @staticmethod
    def markdown_to_word(markdown_path: str, output_path: Optional[str] = None, template_path: Optional[str] = None):
        """
        将 Markdown 文件转换为 Word 文档
        
        参数：
            markdown_path: Markdown 文件路径
            output_path: 输出 Word 文件路径（可选）
            template_path: Word 模板文件路径（可选）
        
        返回：
            output_path: 输出文件路径
        """
        md_path = Path(markdown_path)
        if not md_path.exists():
            raise FileNotFoundError(f"Markdown 文件不存在: {markdown_path}")
        
        # 确定输出路径
        if not output_path:
            output_path = md_path.with_suffix('.docx')
        
        # 创建导出器
        exporter = GeoWordExporter(template_path)
        
        # 读取 Markdown 内容
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # 统一换行符
        content = content.replace('\r\n', '\n').replace('\r', '\n')

        # 简单解析 Markdown 并添加到文档
        lines = content.split('\n')
        current_section = ""
        table_lines = []

        for line in lines:
            # 处理标题
            if line.startswith('# '):
                if table_lines:
                    exporter.add_table_from_markdown('\n'.join(table_lines))
                    table_lines = []
                if current_section:
                    exporter.add_paragraph(current_section.strip())
                    current_section = ""
                exporter.add_section(line[2:], level=1)
            elif line.startswith('## '):
                if table_lines:
                    exporter.add_table_from_markdown('\n'.join(table_lines))
                    table_lines = []
                if current_section:
                    exporter.add_paragraph(current_section.strip())
                    current_section = ""
                exporter.add_section(line[3:], level=2)
            elif line.startswith('### '):
                if table_lines:
                    exporter.add_table_from_markdown('\n'.join(table_lines))
                    table_lines = []
                if current_section:
                    exporter.add_paragraph(current_section.strip())
                    current_section = ""
                exporter.add_section(line[4:], level=3)
            # 收集表格行
            elif line.startswith('|'):
                if current_section:
                    exporter.add_paragraph(current_section.strip())
                    current_section = ""
                table_lines.append(line)
            # 处理段落
            elif line.strip():
                if table_lines:
                    exporter.add_table_from_markdown('\n'.join(table_lines))
                    table_lines = []
                current_section += line + '\n'
            # 处理空行
            elif current_section:
                exporter.add_paragraph(current_section.strip())
                current_section = ""

        # 处理末尾残余的表格或段落
        if table_lines:
            exporter.add_table_from_markdown('\n'.join(table_lines))
        elif current_section:
            exporter.add_paragraph(current_section.strip())
        
        # 保存文档
        exporter.save(output_path)
        return str(output_path)
    
    def save(self, output_path: str):
        """
        保存 Word 文档

        参数：
            output_path: 输出文件路径
        """
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(output)
        print(f"[OK] Word 文档已保存: {output}")


# 便捷函数
def export_markdown_to_word(markdown_path: str, output_path: Optional[str] = None, template_path: Optional[str] = None):
    """
    便捷函数：将 Markdown 文件转换为 Word 文档
    
    参数：
        markdown_path: Markdown 文件路径
        output_path: 输出 Word 文件路径（可选）
        template_path: Word 模板文件路径（可选）
    
    返回：
        output_path: 输出文件路径
    """
    return GeoWordExporter.markdown_to_word(markdown_path, output_path, template_path)


if __name__ == "__main__":
    # 测试示例
    print("Word 导出器测试...")
    print("\n使用方法：")
    print("  from geo_word_exporter import GeoWordExporter")
    print("  exporter = GeoWordExporter(template_path='模板.docx')")
    print("  exporter.add_section('第一章', '这是内容', level=1)")
    print("  exporter.save('输出.docx')")
    print("\n或使用便捷函数：")
    print("  export_markdown_to_word('报告.md', '输出.docx')")
